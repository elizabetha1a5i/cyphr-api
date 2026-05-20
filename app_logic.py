"""
Cyphr template generator — populates SOW and Proposal from project data
Usage: python generate_templates.py project_data.json
"""
import json, sys, os, copy, re, shutil
from pathlib import Path

# ── DOCX SOW ──────────────────────────────────────────────────────────────
def generate_sow(data, output_path):
    from docx import Document
    from docx.shared import Pt
    from datetime import datetime
    
    doc = Document('/mnt/user-data/uploads/Blue_Square_x_Cyphr_SOW_TEMPLATE.docx')
    
    client = data.get('clientName', 'CLIENT')
    project = data.get('projectName', 'PROJECT')
    budget = data.get('budget', '0')
    timeline = data.get('timeline', 'TBC')
    brief = data.get('briefOutput', '')
    estimate = data.get('estimateOutput', '')
    sow_content = data.get('sowOutput', '')
    today = datetime.today().strftime('%-d %B %Y')
    
    # Replace in opening paragraph (effective date + parties)
    for para in doc.paragraphs:
        if 'Statement of Work' in para.text and 'effective from' in para.text:
            for run in para.runs:
                run.text = run.text.replace('20th March 2026', today)
        if 'Blue Square Marketing Limited' in para.text and 'agreement is entered' in para.text:
            for run in para.runs:
                run.text = run.text.replace('Blue Square Marketing Limited', client)

    # Populate the main table
    table = doc.tables[0]
    
    # Parse SOW content into sections
    sections = parse_sow_sections(sow_content, brief, estimate, client, project, budget, timeline)
    
    # Map sections to table rows
    row_map = {
        3: sections.get('summary', f'Cyphr will {brief[:300] if brief else "deliver the agreed project scope."}'),
        5: sections.get('objectives', estimate[:400] if estimate else 'Deliverables to be confirmed per project scope.'),
        7: sections.get('assumptions', 'Client to provide all required content and access within agreed timelines.\nAll third-party integrations and API access to be arranged by client prior to project kick-off.'),
        9: sections.get('responsibilities', f'Cyphr will be responsible for all design, build and delivery activities.\n{client} will be responsible for content provision, stakeholder sign-off and UAT feedback.'),
        11: 'United Kingdom',
        15: f'Cyphr will meet with the {client} team regularly to discuss requirements and progress. Weekly status updates will be provided throughout the project.',
        17: f'Cyphr: Verity Smout\n{client}: TBC',
        19: f'Cyphr will provide regular project updates during delivery. A shared project tracker will be maintained throughout.',
        23: 'Cyphr will address critical issues within 24 hours. All bugs will be tracked and resolved within agreed SLAs.',
        26: sections.get('fee', f'Services will be charged on a fixed price basis of £{format_budget(budget)} as set out in the agreed estimate.'),
        28: 'Services will be invoiced at project milestones as agreed. Payment terms: 30 days from invoice date.',
        30: f'Invoice to be submitted to Accounts Payable | {client}',
        32: 'A change request (CR) will result in an incremental cost estimate. All CRs require written approval before work commences.',
    }
    
    for row_idx, content in row_map.items():
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            # Update both cells (merged row workaround)
            for cell in row.cells[:1]:
                # Clear existing paragraphs
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                # Set first paragraph
                if cell.paragraphs:
                    cell.paragraphs[0].runs[0].text = content if cell.paragraphs[0].runs else ''
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run(content)

    # Update milestone table row (row 21) with timeline data
    if len(table.rows) > 21:
        milestone_row = table.rows[21]
        milestone_text = sections.get('milestones', f'Project kick-off: Week 1\nDiscovery complete: Week 3\nDesign sign-off: Week 5\nBuild complete: Week {get_weeks(timeline)}\nUAT: Week {get_weeks(timeline)+1}\nLaunch: Week {get_weeks(timeline)+2}')
        for cell in milestone_row.cells[:1]:
            if cell.paragraphs:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = milestone_text
                else:
                    cell.paragraphs[0].add_run(milestone_text)

    doc.save(output_path)
    print(f"SOW saved: {output_path}")


def parse_sow_sections(sow_text, brief, estimate, client, project, budget, timeline):
    """Parse AI-generated SOW text into sections"""
    sections = {}
    
    if not sow_text:
        sections['summary'] = f'{client} requires Cyphr to deliver {project}. This engagement covers the full project lifecycle from discovery through to delivery.'
        sections['objectives'] = f'Deliver {project} on time and within budget of £{format_budget(budget)}.'
        sections['fee'] = f'Fixed price of £{format_budget(budget)} for the full scope as described.'
        return sections
    
    # Try to extract sections from structured SOW text
    lines = sow_text.split('\n')
    current_section = None
    section_text = []
    
    section_map = {
        '1.1': 'summary', 'project summary': 'summary',
        '1.2': 'objectives', 'objectives': 'objectives',
        '1.3': 'assumptions', 'assumptions': 'assumptions',
        '4.1': 'milestones', 'milestones': 'milestones',
        '5.1': 'fee', 'fee': 'fee',
    }
    
    for line in lines:
        lower = line.lower().strip()
        matched = None
        for key, sec in section_map.items():
            if key in lower:
                if current_section and section_text:
                    sections[current_section] = '\n'.join(section_text).strip()
                current_section = sec
                section_text = []
                matched = True
                break
        if not matched and current_section and line.strip():
            section_text.append(line.strip())
    
    if current_section and section_text:
        sections[current_section] = '\n'.join(section_text).strip()
    
    # Fallbacks
    if 'summary' not in sections:
        sections['summary'] = sow_text[:300].strip()
    if 'fee' not in sections:
        sections['fee'] = f'Fixed price of £{format_budget(budget)} as per the agreed estimate.'
    
    return sections


def format_budget(budget):
    try:
        return f"{int(float(budget)):,}"
    except:
        return str(budget)

def get_weeks(timeline):
    """Extract number of weeks from timeline string"""
    match = re.search(r'(\d+)', str(timeline))
    return int(match.group(1)) if match else 10


# ── PPTX PROPOSAL ─────────────────────────────────────────────────────────
def generate_proposal(data, output_path):
    import subprocess, os
    
    client = data.get('clientName', 'CLIENT')
    project = data.get('projectName', 'PROJECT')
    budget = data.get('budget', '0')
    timeline = data.get('timeline', 'TBC')
    brief = data.get('briefOutput', '')
    estimate = data.get('estimateOutput', '')
    proposal = data.get('proposalOutput', '')
    
    # Work from unpacked template
    src = '/mnt/user-data/uploads/Cyphr_x_Blue_Square____Cost_Proposal_TEMPLATE_____INTERNAL_SHARE_.pptx'
    work_dir = '/home/claude/proposal_work/'
    
    if os.path.exists(work_dir):
        shutil.rmtree(work_dir)
    
    # Unpack
    result = subprocess.run(
        ['python3', '/mnt/skills/public/pptx/scripts/office/unpack.py', src, work_dir],
        capture_output=True, text=True
    )
    
    slides_dir = os.path.join(work_dir, 'ppt/slides/')
    
    # Parse proposal content
    sections = parse_proposal_sections(proposal, brief, estimate, client, project, budget, timeline)
    
    # Slide 1: Cover — replace CLIENT and PROJECT NAME
    replace_in_slide(slides_dir + 'slide1.xml', {
        'CLIENT': client.upper(),
        'PROJECT NAME': project.upper() if project else 'PROPOSAL',
        'Cost Estimate': 'Commercial Proposal',
        'CYPHR \nX BLUE SQUARE': f'CYPHR X {client.upper()}',
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
    })
    
    # Slide 2: Executive Summary
    replace_in_slide(slides_dir + 'slide2.xml', {
        'This proposal outlines….': sections.get('exec_summary', f'This proposal outlines Cyphr\'s approach to delivering {project} for {client}.'),
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
    })
    
    # Slide 3: The Ask
    replace_in_slide(slides_dir + 'slide3.xml', {
        'Activity Overview….': sections.get('the_ask', f'{client} needs a partner to deliver {project}.'),
        'The Digital User Journey': sections.get('approach_headline', 'Our Approach'),
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
    })
    
    # Slide 6: Milestones (keep structure, update footer)
    replace_in_slide(slides_dir + 'slide6.xml', {
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
    })
    
    # Slide 8: Cost breakdown
    cost_section = sections.get('cost', '')
    replace_in_slide(slides_dir + 'slide8.xml', {
        'Core Roadshow Experience Web App Design &amp; Build': project,
        'Core Roadshow Experience Web App Design & Build': project,
        '£25,314': f'£{format_budget(budget)}',
        'TOTAL': 'TOTAL',
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
        '2 weeks + Tour duration adhoc support': timeline,
    })
    
    # Slide 9: Thank you
    replace_in_slide(slides_dir + 'slide9.xml', {
        'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
    })
    
    # Update all slides footers
    for i in range(1, 10):
        slide_path = slides_dir + f'slide{i}.xml'
        if os.path.exists(slide_path):
            replace_in_slide(slide_path, {
                'CYPHR X BLUE SQUARE': f'CYPHR X {client.upper()}',
                'PRESENTATION': 'PROPOSAL',
            })
    
    # Clean and pack
    subprocess.run(['python3', '/mnt/skills/public/pptx/scripts/clean.py', work_dir], capture_output=True)
    subprocess.run([
        'python3', '/mnt/skills/public/pptx/scripts/office/pack.py',
        work_dir, output_path, '--original', src
    ], capture_output=True)
    
    print(f"Proposal saved: {output_path}")


def replace_in_slide(slide_path, replacements):
    """Replace text in a slide XML file"""
    if not os.path.exists(slide_path):
        return
    with open(slide_path, 'r', encoding='utf-8') as f:
        content = f.read()
    for old, new in replacements.items():
        content = content.replace(old, new)
    with open(slide_path, 'w', encoding='utf-8') as f:
        f.write(content)


def parse_proposal_sections(proposal_text, brief, estimate, client, project, budget, timeline):
    sections = {}
    
    if not proposal_text:
        sections['exec_summary'] = f'Cyphr proposes to deliver {project} for {client} within {timeline} at a total investment of £{format_budget(budget)}. This proposal outlines our approach, deliverables, team and commercial terms.'
        sections['the_ask'] = brief[:300] if brief else f'{client} requires a digital solution to address their key business objectives.'
        sections['approach_headline'] = 'Our Approach'
        sections['cost'] = estimate[:200] if estimate else f'Total fixed price: £{format_budget(budget)}'
        return sections
    
    lines = proposal_text.split('\n')
    current = None
    buf = []
    
    section_triggers = {
        'executive summary': 'exec_summary',
        '1.': 'exec_summary',
        'the challenge': 'the_ask',
        '2.': 'the_ask',
        'our approach': 'approach_headline',
        '3.': 'approach_headline',
        'what we': 'deliverables',
        '4.': 'deliverables',
        'investment': 'cost',
        '7.': 'cost',
    }
    
    for line in lines:
        lower = line.lower().strip()
        matched = False
        for trigger, sec in section_triggers.items():
            if lower.startswith(trigger) or (trigger.endswith('.') and lower.startswith(trigger)):
                if current and buf:
                    sections[current] = '\n'.join(buf).strip()
                current = sec
                buf = []
                matched = True
                break
        if not matched and current and line.strip():
            buf.append(line.strip())
    
    if current and buf:
        sections[current] = '\n'.join(buf).strip()
    
    # Fallbacks
    if 'exec_summary' not in sections:
        # Use first 3 paragraphs of proposal
        paras = [p.strip() for p in proposal_text.split('\n\n') if p.strip()]
        sections['exec_summary'] = paras[0][:400] if paras else f'Cyphr proposes to deliver {project} for {client}.'
    
    return sections


# ── XLSX ESTIMATE (enhanced version) ──────────────────────────────────────
def generate_estimate(data, output_path):
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    # Copy template
    shutil.copy(
        '/mnt/user-data/uploads/Copy_of_TEMPLATE_ESTIMATE_2026____CLIENT_NAME____PROJECT_NAME_____DATE__STATUS__SHARE_INTERNAL_.xlsx',
        output_path
    )
    
    wb = openpyxl.load_workbook(output_path)
    ws = wb['TEMPLATE']
    
    client = data.get('clientName', 'CLIENT')
    project = data.get('projectName', 'PROJECT')
    budget = data.get('budget', '0')
    timeline = data.get('timeline', 'TBC')
    total_days = int(float(data.get('days', 40) or 40))
    
    from datetime import datetime
    today = datetime.today().strftime('%d/%m/%Y')
    
    # Update header row A1 if it exists
    ws['A1'] = f'Cyphr Cost Estimate'
    
    # Find and replace CLIENT_NAME and PROJECT_NAME placeholders
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cell.value = (cell.value
                    .replace('CLIENT_NAME', client)
                    .replace('PROJECT_NAME', project)
                    .replace('CLIENT NAME', client)
                    .replace('PROJECT NAME', project)
                    .replace('DATE', today)
                    .replace('SHARE_INTERNAL', 'INTERNAL')
                )
    
    wb.save(output_path)
    print(f"Estimate saved: {output_path}")


# ── XLSX GANTT (enhanced version) ─────────────────────────────────────────
def generate_gantt(data, output_path):
    import openpyxl
    
    # Copy template
    shutil.copy(
        '/mnt/user-data/uploads/Cyphr_x_CLIENT____PROJECT____Gantt_Chart_TEMPLATE.xlsx',
        output_path
    )
    
    wb = openpyxl.load_workbook(output_path)
    ws = wb.active
    
    client = data.get('clientName', 'CLIENT')
    project = data.get('projectName', 'PROJECT')
    
    # Replace placeholders
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                cell.value = (cell.value
                    .replace('CLIENT', client)
                    .replace('PROJECT', project)
                )
    
    wb.save(output_path)
    print(f"Gantt saved: {output_path}")


# ── MAIN ───────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python generate_templates.py project_data.json [output_dir]")
        sys.exit(1)
    
    with open(sys.argv[1]) as f:
        data = json.load(f)
    
    out_dir = sys.argv[2] if len(sys.argv) > 2 else '/home/claude/output'
    os.makedirs(out_dir, exist_ok=True)
    
    slug = data.get('clientName', 'project').lower().replace(' ', '_')
    
    generate_sow(data, f'{out_dir}/{slug}_sow.docx')
    generate_proposal(data, f'{out_dir}/{slug}_proposal.pptx')
    generate_estimate(data, f'{out_dir}/{slug}_estimate.xlsx')
    generate_gantt(data, f'{out_dir}/{slug}_gantt.xlsx')
    
    print(f"\nAll files generated in {out_dir}/")
